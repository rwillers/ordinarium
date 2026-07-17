import io

from .docx_blocks import extract_blocks, extract_inline_tokens

PRINT_RUBRIC_COLOR = "C62F7C"


def render_docx_bytes(context):
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("DOCX export requires python-docx to be installed.") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal_style = styles["Normal"]
    normal_style.font.name = "Georgia"
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(10)

    def ensure_style(name, base_style="Normal"):
        try:
            return styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            if base_style:
                style.base_style = styles[base_style]
            return style

    title_style = ensure_style("LiturgyTitle")
    title_style.font.size = Pt(15)
    title_style.font.bold = False
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_before = Pt(7)
    title_style.paragraph_format.space_after = Pt(7)

    subtitle_style = ensure_style("LiturgySubtitle")
    subtitle_style.font.size = Pt(12)
    subtitle_style.font.italic = True
    subtitle_style.font.color.rgb = RGBColor.from_string(PRINT_RUBRIC_COLOR)
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(9)

    section_style = ensure_style("LiturgySection")
    section_style.font.size = Pt(10)
    section_style.font.color.rgb = RGBColor.from_string(PRINT_RUBRIC_COLOR)
    section_style.font.all_caps = True
    section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    section_style.paragraph_format.space_before = Pt(11)
    section_style.paragraph_format.space_after = Pt(11)
    section_style.paragraph_format.keep_with_next = True

    body_style = ensure_style("LiturgyBody")
    body_style.font.size = Pt(11)
    body_style.paragraph_format.space_before = Pt(0)
    body_style.paragraph_format.space_after = Pt(10)

    pre_style = ensure_style("LiturgyPre", base_style="LiturgyBody")
    pre_style.paragraph_format.space_before = Pt(0)
    pre_style.paragraph_format.space_after = Pt(10)
    pre_style.paragraph_format.keep_together = True

    scripture_style = ensure_style("LiturgyScripture", base_style="LiturgyBody")
    scripture_style.font.size = Pt(9)
    scripture_style.font.color.rgb = RGBColor.from_string(PRINT_RUBRIC_COLOR)
    scripture_style.font.all_caps = True
    scripture_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    scripture_style.paragraph_format.space_before = Pt(0)
    scripture_style.paragraph_format.space_after = Pt(10)

    footer_style = ensure_style("LiturgyFooter")
    footer_style.font.size = Pt(9)
    footer_style.font.color.rgb = RGBColor.from_string(PRINT_RUBRIC_COLOR)
    footer_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_style.paragraph_format.space_before = Pt(18)
    footer_style.paragraph_format.space_after = Pt(0)

    bullet_style = ensure_style("LiturgyBullet", base_style="LiturgyBody")
    bullet_style.paragraph_format.left_indent = Pt(18)
    bullet_style.paragraph_format.first_line_indent = Pt(-10)
    rubric_color = RGBColor.from_string(PRINT_RUBRIC_COLOR)

    def add_inline_tokens(paragraph, tokens):
        for token in tokens:
            if token["kind"] == "break":
                paragraph.add_run().add_break()
                continue
            run = paragraph.add_run(token["text"])
            run.bold = token["bold"]
            run.italic = token["italic"]
            if token["rubric"]:
                run.font.color.rgb = rubric_color

    title_paragraph = document.add_paragraph(style=title_style)
    title_paragraph.add_run(context.get("title", ""))
    title_paragraph.add_run().add_break()
    rite_run = title_paragraph.add_run(context.get("rite", ""))
    rite_run.italic = True
    rite_run.font.size = Pt(8)
    rite_run.font.color.rgb = rubric_color

    if context.get("service_title"):
        document.add_paragraph(context["service_title"], style=subtitle_style)

    for ordinary in context["ordinaries"]:
        if ordinary.get("show_title") and ordinary.get("title_inline_html"):
            paragraph = document.add_paragraph(style=section_style)
            add_inline_tokens(
                paragraph, extract_inline_tokens(ordinary["title_inline_html"])
            )
        blocks = extract_blocks(
            ordinary.get("body_html", ""),
            treat_ul_as_variants=ordinary.get("type") != "custom",
        )
        _add_blocks(
            document,
            blocks,
            add_inline_tokens,
            body_style,
            pre_style,
            scripture_style,
            bullet_style,
        )

    footer = document.add_paragraph(style=footer_style)
    footer.add_run(f"{context.get('title', '')} - {context.get('rite', '')}")
    if context.get("service_title"):
        footer.add_run().add_break()
        footer.add_run(context["service_title"])
    if context.get("service_date_display"):
        footer.add_run().add_break()
        footer.add_run(f"({context['service_date_display']})")
    footer.add_run().add_break()
    footer.add_run(f"Generated as of {context.get('generated_at_display', '')}")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _add_blocks(
    document,
    blocks,
    add_inline_tokens,
    body_style,
    pre_style,
    scripture_style,
    bullet_style,
):
    for block in blocks:
        if block["kind"] == "paragraph":
            style = scripture_style if block["variant"] == "scripture" else body_style
            paragraph = document.add_paragraph(style=style)
            add_inline_tokens(paragraph, block["tokens"])
            continue
        if block["kind"] == "pre":
            paragraph = document.add_paragraph(style=pre_style)
            paragraph.add_run(block["text"]).bold = True
            continue
        if block["kind"] == "bullet":
            paragraph = document.add_paragraph(style=bullet_style)
            paragraph.add_run("\u2022 ")
            add_inline_tokens(paragraph, block["tokens"])
