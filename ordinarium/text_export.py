import io
import re
from html.parser import HTMLParser

from flask import current_app, render_template

from .text_rendering import build_rendered_ordinaries

PRINT_RUBRIC_COLOR = "C62F7C"

EXPORT_TEXT_CSS = """
@page {
    margin: 1in;
}

html, body {
    margin: 0;
    padding: 0;
}

body {
    background: #fff;
    color: #000;
    font-family: Georgia, serif;
    font-size: 11pt;
}

#text {
    padding-top: 0;
}

#text h1 {
    font-size: 1.36em;
    font-weight: 200;
    text-align: center;
    text-transform: uppercase;
    letter-spacing: 0.2em;
    line-height: 1.67em;
    margin: 0.67em 0;
}

#text h1 small {
    display: inline-block;
    font-size: 0.8em;
    letter-spacing: 0.1em;
    color: #c62f7c;
}

#text h1 small em {
    font-size: 0.67em;
    text-transform: none;
}

#text h2 {
    font-size: 1.1em;
    font-weight: 400;
    font-style: italic;
    text-align: center;
    color: #c62f7c;
    margin: 0.83em 0;
}

#text h3 {
    font-size: 0.91em;
    font-weight: 400;
    letter-spacing: 0.1em;
    text-align: center;
    text-transform: uppercase;
    color: #c62f7c;
    margin: 1em 0;
    break-after: avoid-page;
}

#text h6 {
    font-size: 0.82em;
    font-weight: 400;
    text-align: right;
    text-transform: uppercase;
    color: #c62f7c;
    margin: -1.11em 0 1.11em 0;
}

#text strong {
    font-weight: 600;
}

#text p em,
#text li em,
#text pre em,
#text code em {
    font-style: italic;
    color: #c62f7c;
}

#text small {
    font-size: 0.81em;
}

#text p {
    margin: 0.91em 0;
    break-inside: avoid;
}

#text pre {
    margin: 0 auto;
    white-space: pre-wrap;
    overflow-wrap: anywhere;
    break-inside: avoid;
}

#text code {
    color: #000;
    font-size: 1em;
    font-family: Georgia, serif;
    font-weight: 600;
}

#text .text-element:not(.text-element-custom) ul {
    margin: 0;
    padding: 0;
    list-style: none;
}

#text .variant-list {
    margin: 0.2em 0 0.9em;
}

#text .variant-option p {
    margin: 0.5em 0;
}

#text .variant-or {
    margin: 0.25em 0;
    text-align: center;
}

#text .variant-or em {
    color: #c62f7c;
    font-style: italic;
}

#text .text-element.text-element-custom ul {
    margin: 0.5em 0 0.9em 1.4em;
    padding: 0;
    list-style: disc;
}

#text .text-element.text-element-custom ul li {
    display: list-item;
    white-space: normal;
}

#text .text-element.text-element-custom ul li + li {
    margin-top: 0.25em;
}

#text p > em:first-child:not(:last-child) {
    display: inline-block;
    margin-right: 0.5em;
    min-width: 4.5em;
    font-size: 0.91em;
    text-align: right;
}

#text p > .trailing-indent {
    display: inline-block;
    width: calc(100% - 5em);
    vertical-align: top;
}

#text .text-footer {
    font-size: 0.8em;
    text-align: center;
    margin-top: 2em;
    color: #c62f7c;
}
"""


def build_text_export_context(service_id, saved_service, saved_data, user_id=None):
    payload = build_rendered_ordinaries(
        service_id,
        saved_service,
        saved_data,
        user_id=user_id,
        include_metadata=True,
    )
    if not payload:
        return None

    markdown_filter = current_app.jinja_env.filters.get("markdown")
    markdown_user_filter = current_app.jinja_env.filters.get("markdown_user")
    trailing_indent_filter = current_app.jinja_env.filters.get("trailing_indent")

    if not markdown_filter or not markdown_user_filter or not trailing_indent_filter:
        raise RuntimeError("Required markdown filters are not configured.")

    ordinaries = []
    previous_title = None
    for ordinary in payload["ordinaries"]:
        render_markdown = (
            markdown_user_filter
            if ordinary.get("type") == "custom"
            else markdown_filter
        )
        title_markdown = ordinary.get("title") or ""
        title_html = str(render_markdown(title_markdown))
        title_inline_html = _strip_wrapping_paragraph(title_html)

        body_markdown = ordinary.get("text") or ""
        body_html = str(render_markdown(body_markdown))
        body_html = str(trailing_indent_filter(body_html))
        if ordinary.get("type") != "custom":
            body_html = _rewrite_variant_lists(body_html)

        show_title = bool(title_markdown) and title_markdown != previous_title
        previous_title = title_markdown

        ordinaries.append(
            {
                "type": ordinary.get("type"),
                "title_markdown": title_markdown,
                "title_inline_html": title_inline_html,
                "body_html": body_html,
                "show_title": show_title,
            }
        )

    payload["service_id"] = service_id
    payload["ordinaries"] = ordinaries
    return payload


def render_text_export_html(context):
    return render_template("export_text.html", export_css=EXPORT_TEXT_CSS, **context)


def render_pdf_bytes(html_text, base_url=None):
    try:
        from weasyprint import HTML
    except ImportError as exc:
        raise RuntimeError("PDF export requires weasyprint to be installed.") from exc
    try:
        return HTML(string=html_text, base_url=base_url).write_pdf()
    except Exception as exc:
        raise RuntimeError(
            "PDF export failed in WeasyPrint runtime. Verify server libraries for "
            "Pango, Cairo, and Fontconfig are installed."
        ) from exc


def render_docx_bytes(context):
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("DOCX export requires python-docx to be installed.") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal_style = styles["Normal"]
    normal_style.font.name = "Georgia"
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(10)

    def ensure_style(name, base_style="Normal"):
        try:
            return styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            if base_style:
                style.base_style = styles[base_style]
            return style

    title_style = ensure_style("LiturgyTitle")
    title_style.font.size = Pt(15)
    title_style.font.bold = False
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_before = Pt(7)
    title_style.paragraph_format.space_after = Pt(7)

    subtitle_style = ensure_style("LiturgySubtitle")
    subtitle_style.font.size = Pt(12)
    subtitle_style.font.italic = True
    subtitle_style.font.color.rgb = RGBColor.from_string(PRINT_RUBRIC_COLOR)
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(9)

    section_style = ensure_style("LiturgySection")
    section_style.font.size = Pt(10)
    section_style.font.color.rgb = RGBColor.from_string(PRINT_RUBRIC_COLOR)
    section_style.font.all_caps = True
    section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    section_style.paragraph_format.space_before = Pt(11)
    section_style.paragraph_format.space_after = Pt(11)
    section_style.paragraph_format.keep_with_next = True

    body_style = ensure_style("LiturgyBody")
    body_style.font.size = Pt(11)
    body_style.paragraph_format.space_before = Pt(0)
    body_style.paragraph_format.space_after = Pt(10)

    pre_style = ensure_style("LiturgyPre", base_style="LiturgyBody")
    pre_style.paragraph_format.space_before = Pt(0)
    pre_style.paragraph_format.space_after = Pt(10)
    pre_style.paragraph_format.keep_together = True

    scripture_style = ensure_style("LiturgyScripture", base_style="LiturgyBody")
    scripture_style.font.size = Pt(9)
    scripture_style.font.color.rgb = RGBColor.from_string(PRINT_RUBRIC_COLOR)
    scripture_style.font.all_caps = True
    scripture_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    scripture_style.paragraph_format.space_before = Pt(0)
    scripture_style.paragraph_format.space_after = Pt(10)

    footer_style = ensure_style("LiturgyFooter")
    footer_style.font.size = Pt(9)
    footer_style.font.color.rgb = RGBColor.from_string(PRINT_RUBRIC_COLOR)
    footer_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_style.paragraph_format.space_before = Pt(18)
    footer_style.paragraph_format.space_after = Pt(0)

    bullet_style = ensure_style("LiturgyBullet", base_style="LiturgyBody")
    bullet_style.paragraph_format.left_indent = Pt(18)
    bullet_style.paragraph_format.first_line_indent = Pt(-10)

    rubric_color = RGBColor.from_string(PRINT_RUBRIC_COLOR)

    def add_inline_tokens(paragraph, tokens):
        for token in tokens:
            if token["kind"] == "break":
                paragraph.add_run().add_break()
                continue
            run = paragraph.add_run(token["text"])
            if token["bold"]:
                run.bold = True
            if token["italic"]:
                run.italic = True
            if token["rubric"]:
                run.font.color.rgb = rubric_color

    title_paragraph = document.add_paragraph(style=title_style)
    title_paragraph.add_run(context.get("title", ""))
    title_paragraph.add_run().add_break()
    rite_run = title_paragraph.add_run(context.get("rite", ""))
    rite_run.italic = True
    rite_run.font.size = Pt(8)
    rite_run.font.color.rgb = rubric_color

    if context.get("service_title"):
        document.add_paragraph(context["service_title"], style=subtitle_style)

    for ordinary in context["ordinaries"]:
        if ordinary.get("show_title") and ordinary.get("title_inline_html"):
            section = document.add_paragraph(style=section_style)
            add_inline_tokens(
                section, _extract_inline_tokens(ordinary.get("title_inline_html", ""))
            )
        for block in _extract_blocks(
            ordinary.get("body_html", ""),
            treat_ul_as_variants=ordinary.get("type") != "custom",
        ):
            if block["kind"] == "paragraph":
                style_name = (
                    scripture_style if block["variant"] == "scripture" else body_style
                )
                paragraph = document.add_paragraph(style=style_name)
                add_inline_tokens(paragraph, block["tokens"])
                continue
            if block["kind"] == "pre":
                paragraph = document.add_paragraph(style=pre_style)
                run = paragraph.add_run(block["text"])
                run.bold = True
                continue
            if block["kind"] == "bullet":
                paragraph = document.add_paragraph(style=bullet_style)
                paragraph.add_run("\u2022 ")
                add_inline_tokens(paragraph, block["tokens"])

    footer = document.add_paragraph(style=footer_style)
    footer.add_run(f"{context.get('title', '')} - {context.get('rite', '')}")
    if context.get("service_title"):
        footer.add_run().add_break()
        footer.add_run(context["service_title"])
    if context.get("service_date_display"):
        footer.add_run().add_break()
        footer.add_run(f"({context['service_date_display']})")
    footer.add_run().add_break()
    footer.add_run(f"Generated as of {context.get('generated_at_display', '')}")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_export_filename(context, extension):
    base_parts = []
    service_date = context.get("service_date")
    if service_date:
        base_parts.append(service_date)
    service_title = context.get("service_title") or context.get("title")
    if service_title:
        base_parts.append(service_title)
    base = " ".join(base_parts).strip() or f"service-{context.get('service_id', '')}"
    slug = re.sub(r"[^a-z0-9]+", "-", base.lower()).strip("-")
    if not slug:
        slug = f"service-{context.get('service_id', '')}"
    return f"{slug}.{extension}"


def _strip_wrapping_paragraph(value):
    if not value:
        return ""
    match = re.fullmatch(r"\s*<p\b[^>]*>(.*)</p>\s*", value, flags=re.DOTALL)
    if not match:
        return value.strip()
    inner = match.group(1).strip()
    if "<p" in inner.lower():
        return value.strip()
    return inner


def _extract_inline_tokens(fragment):
    root = _parse_fragment(fragment)
    return _trim_edge_breaks(_collect_inline_tokens(root))


def _extract_blocks(fragment, treat_ul_as_variants=True):
    root = _parse_fragment(fragment)
    blocks = []
    for child in root["children"]:
        if child["type"] == "text":
            if child["text"].strip():
                blocks.append(
                    {
                        "kind": "paragraph",
                        "variant": "body",
                        "tokens": _trim_edge_breaks(
                            [
                                {
                                    "kind": "text",
                                    "text": child["text"],
                                    **_base_text_style(),
                                }
                            ]
                        ),
                    }
                )
            continue

        tag = child["tag"]
        if tag == "p":
            tokens = _trim_edge_breaks(_collect_inline_tokens(child))
            if _tokens_have_content(tokens):
                blocks.append(
                    {"kind": "paragraph", "variant": "body", "tokens": tokens}
                )
            continue
        if tag == "h6":
            tokens = _trim_edge_breaks(_collect_inline_tokens(child))
            if _tokens_have_content(tokens):
                blocks.append(
                    {"kind": "paragraph", "variant": "scripture", "tokens": tokens}
                )
            continue
        if tag == "pre":
            text = _extract_pre_text(child)
            if text.strip():
                blocks.append({"kind": "pre", "text": text})
            continue
        if tag == "ul":
            items = []
            for list_item in child["children"]:
                if list_item["type"] != "element" or list_item["tag"] != "li":
                    continue
                tokens = _trim_edge_breaks(
                    _collect_inline_tokens(list_item, parent_tag="li")
                )
                if _tokens_have_content(tokens):
                    items.append(tokens)
            if not items:
                continue
            if treat_ul_as_variants:
                for index, tokens in enumerate(items):
                    blocks.append(
                        {"kind": "paragraph", "variant": "body", "tokens": tokens}
                    )
                    if index < len(items) - 1:
                        blocks.append(
                            {
                                "kind": "paragraph",
                                "variant": "body",
                                "tokens": [
                                    {
                                        "kind": "text",
                                        "text": "Or",
                                        "bold": False,
                                        "italic": True,
                                        "rubric": True,
                                    }
                                ],
                            }
                        )
                continue
            for tokens in items:
                blocks.append({"kind": "bullet", "tokens": tokens})
            continue
        blocks.extend(
            _extract_blocks(
                _serialize_children(child),
                treat_ul_as_variants=treat_ul_as_variants,
            )
        )
    return blocks


def _rewrite_variant_lists(fragment):
    root = _parse_fragment(fragment)
    return "".join(
        _serialize_node_with_variant_lists(child) for child in root["children"]
    )


def _serialize_node_with_variant_lists(node):
    if node["type"] == "text":
        return node["text"]

    tag = node["tag"]
    if tag == "br":
        return "<br>"
    if tag == "ul":
        options = [
            child
            for child in node["children"]
            if child["type"] == "element" and child["tag"] == "li"
        ]
        if not options:
            return ""
        parts = ['<div class="variant-list">']
        for index, option in enumerate(options):
            parts.append('<div class="variant-option">')
            parts.extend(
                _serialize_node_with_variant_lists(child)
                for child in option["children"]
            )
            parts.append("</div>")
            if index < len(options) - 1:
                parts.append('<p class="variant-or"><em>Or</em></p>')
        parts.append("</div>")
        return "".join(parts)

    attrs = "".join(f' {key}="{value}"' for key, value in node["attrs"].items())
    inner = "".join(
        _serialize_node_with_variant_lists(child) for child in node["children"]
    )
    return f"<{tag}{attrs}>{inner}</{tag}>"


def _extract_pre_text(node):
    parts = []

    def visit(current):
        for child in current["children"]:
            if child["type"] == "text":
                parts.append(child["text"])
                continue
            if child["tag"] == "br":
                parts.append("\n")
                continue
            visit(child)

    visit(node)
    return "".join(parts).strip("\n")


def _collect_inline_tokens(node, style=None, parent_tag=""):
    if style is None:
        style = _base_text_style()

    tokens = []
    for child in node["children"]:
        if child["type"] == "text":
            tokens.append({"kind": "text", "text": child["text"], **style})
            continue

        tag = child["tag"]
        if tag == "br":
            tokens.append({"kind": "break"})
            continue

        next_style = dict(style)
        if tag in {"strong", "b"}:
            next_style["bold"] = True
        if tag in {"em", "i"}:
            next_style["italic"] = True
            next_style["rubric"] = True
        if tag == "code":
            next_style["bold"] = True

        is_nested_paragraph = tag == "p" and parent_tag == "li"
        if is_nested_paragraph and tokens and tokens[-1]["kind"] != "break":
            tokens.append({"kind": "break"})

        child_tokens = _collect_inline_tokens(child, next_style, parent_tag=tag)
        tokens.extend(child_tokens)

        if is_nested_paragraph and child_tokens:
            if tokens and tokens[-1]["kind"] != "break":
                tokens.append({"kind": "break"})

    return tokens


def _base_text_style():
    return {"bold": False, "italic": False, "rubric": False}


def _tokens_have_content(tokens):
    return any(token["kind"] == "text" and token["text"].strip() for token in tokens)


def _trim_edge_breaks(tokens):
    trimmed = list(tokens)
    while trimmed and trimmed[0]["kind"] == "break":
        trimmed.pop(0)
    while trimmed and trimmed[-1]["kind"] == "break":
        trimmed.pop()
    return trimmed


def _serialize_children(node):
    return "".join(_serialize_node(child) for child in node["children"])


def _serialize_node(node):
    if node["type"] == "text":
        return node["text"]
    if node["tag"] == "br":
        return "<br>"
    attrs = "".join(f' {key}="{value}"' for key, value in node["attrs"].items())
    inner = "".join(_serialize_node(child) for child in node["children"])
    return f"<{node['tag']}{attrs}>{inner}</{node['tag']}>"


def _parse_fragment(fragment):
    parser = _FragmentParser()
    parser.feed(fragment or "")
    parser.close()
    return parser.root


class _FragmentParser(HTMLParser):
    VOID_TAGS = {"br"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = {"type": "element", "tag": "root", "attrs": {}, "children": []}
        self.stack = [self.root]

    def handle_starttag(self, tag, attrs):
        tag_name = tag.lower()
        node = {
            "type": "element",
            "tag": tag_name,
            "attrs": {key: value for key, value in attrs},
            "children": [],
        }
        self.stack[-1]["children"].append(node)
        if tag_name not in self.VOID_TAGS:
            self.stack.append(node)

    def handle_endtag(self, tag):
        tag_name = tag.lower()
        for index in range(len(self.stack) - 1, 0, -1):
            if self.stack[index]["tag"] == tag_name:
                self.stack = self.stack[:index]
                return

    def handle_data(self, data):
        if data:
            self.stack[-1]["children"].append({"type": "text", "text": data})
